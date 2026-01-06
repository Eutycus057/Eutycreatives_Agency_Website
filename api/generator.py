from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.units import inch
import io

def generate_docx(optimized_data):
    doc = Document()
    
    # Set A4 size
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    
    # Set narrow margins for ATS
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Personal Info
    pi = optimized_data.get('personal_info', {})
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run(pi.get('name', '').upper())
    name_run.bold = True
    name_run.font.size = Pt(16)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(pi.get('professional_title', ''))
    title_run.font.size = Pt(12)
    
    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_details = [
        pi.get('email'),
        pi.get('phone'),
        pi.get('location')
    ]
    contact_details.extend(pi.get('links', []))
    contact_info.add_run(" | ".join(filter(None, contact_details)))

    # Professional Summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph(optimized_data.get('professional_summary', ''))

    # Core Skills
    doc.add_heading('CORE SKILLS', level=1)
    doc.add_paragraph(", ".join(optimized_data.get('skills', [])))

    # Work Experience
    doc.add_heading('WORK EXPERIENCE', level=1)
    for job in optimized_data.get('work_experience', []):
        p = doc.add_paragraph()
        p.add_run(f"{job.get('job_title')} – {job.get('company')}").bold = True
        p.add_run(f"\t{job.get('dates')}").italic = True
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5)) 
        
        for bullet in job.get('bullets', []):
            doc.add_paragraph(bullet, style='List Bullet')

    # Education
    doc.add_heading('EDUCATION', level=1)
    for edu in optimized_data.get('education', []):
        p = doc.add_paragraph()
        p.add_run(f"{edu.get('degree')} | {edu.get('institution')}").bold = True
        p.add_run(f"\t{edu.get('dates')}").italic = True
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5))

    # Certifications
    certs = optimized_data.get('certifications', [])
    if certs:
        doc.add_heading('CERTIFICATIONS', level=1)
        for cert in certs:
            doc.add_paragraph(cert, style='List Bullet')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_pdf(optimized_data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    
    # Custom Styles
    name_style = ParagraphStyle('NameStyle', parent=styles['Heading1'], fontSize=18, alignment=1, spaceAfter=2)
    title_style = ParagraphStyle('TitleStyle', parent=styles['Normal'], fontSize=12, alignment=1, spaceAfter=2)
    contact_style = ParagraphStyle('ContactStyle', parent=styles['Normal'], fontSize=10, alignment=1, spaceAfter=10)
    section_title_style = ParagraphStyle('SectionTitle', parent=styles['Heading2'], fontSize=11, spaceBefore=10, spaceAfter=5, textColor="#000000")
    job_title_style = ParagraphStyle('JobTitle', parent=styles['Normal'], fontSize=11, fontWeight='bold', spaceAfter=2)
    bullet_style = ParagraphStyle('Bullet', parent=styles['Normal'], fontSize=10, leftIndent=20, spaceAfter=2)

    story = []

    # Personal Info
    pi = optimized_data.get('personal_info', {})
    story.append(Paragraph(pi.get('name', '').upper(), name_style))
    story.append(Paragraph(pi.get('professional_title', ''), title_style))
    
    contact_info = " | ".join(filter(None, [pi.get('email'), pi.get('phone'), pi.get('location')] + pi.get('links', [])))
    story.append(Paragraph(contact_info, contact_style))

    # Professional Summary
    story.append(Paragraph("PROFESSIONAL SUMMARY", section_title_style))
    story.append(Paragraph(optimized_data.get('professional_summary', ''), styles['Normal']))

    # Skills
    story.append(Paragraph("CORE SKILLS", section_title_style))
    story.append(Paragraph(", ".join(optimized_data.get('skills', [])), styles['Normal']))

    # Work Experience
    story.append(Paragraph("WORK EXPERIENCE", section_title_style))
    for job in optimized_data.get('work_experience', []):
        story.append(Paragraph(f"<b>{job.get('job_title')} – {job.get('company')}</b>", job_title_style))
        story.append(Paragraph(f"<i>{job.get('dates')}</i>", styles['Normal']))
        
        bullets = [ListItem(Paragraph(bullet, bullet_style)) for bullet in job.get('bullets', [])]
        if bullets:
            story.append(ListFlowable(bullets, bulletType='bullet', leftIndent=10))

    # Education
    story.append(Paragraph("EDUCATION", section_title_style))
    for edu in optimized_data.get('education', []):
        story.append(Paragraph(f"<b>{edu.get('degree')} | {edu.get('institution')}</b>", job_title_style))
        story.append(Paragraph(f"<i>{edu.get('dates')}</i>", styles['Normal']))

    # Certifications
    certs = optimized_data.get('certifications', [])
    if certs:
        story.append(Paragraph("CERTIFICATIONS", section_title_style))
        for cert in certs:
            story.append(Paragraph(f"• {cert}", bullet_style))

    doc.build(story)
    buffer.seek(0)
    return buffer
